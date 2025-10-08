"""
DOCUMENT SERVICE
================
Handles document parsing, ingestion, and generation.

FEATURES:
- PDF parsing with LlamaParse (intelligent) + PyPDF2 (fallback)
- DOCX parsing with python-docx
- Document ingestion into Graphiti knowledge graph
- PDF/DOCX/Markdown generation via Pandoc

WHY THIS APPROACH:
- LlamaParse for complex AAOIFI PDFs (better structure extraction)
- PyPDF2 fallback for simple PDFs or when LlamaParse fails
- Graphiti ingestion as EpisodeType.text for searchability
- Pandoc for professional document generation
"""

import logging
import os
from typing import Optional, List, Dict, Any
from pathlib import Path
import aiofiles
from datetime import datetime

# PDF parsing
try:
    from llama_parse import LlamaParse
    LLAMAPARSE_AVAILABLE = True
except ImportError:
    LLAMAPARSE_AVAILABLE = False
    logging.warning("LlamaParse not available, will use PyPDF2 only")

from PyPDF2 import PdfReader

# DOCX parsing
from docx import Document

# Document generation
import pypandoc

# Graphiti integration via MCP
from app.services.graphiti_mcp_service import get_graphiti_mcp_service  # NEW: Using MCP
from app.models import UploadedDocument

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for document parsing, ingestion, and generation."""

    def __init__(
        self,
        llamaparse_api_key: Optional[str] = None,
        upload_dir: str = "./uploads",
        output_dir: str = "./outputs"
    ):
        """
        Initialize DocumentService.

        Args:
            llamaparse_api_key: LlamaParse API key (optional, falls back to PyPDF2)
            upload_dir: Directory for uploaded documents
            output_dir: Directory for generated documents
        """
        self.upload_dir = Path(upload_dir)
        self.output_dir = Path(output_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Initialize LlamaParse if available
        self.llamaparse = None
        if LLAMAPARSE_AVAILABLE and llamaparse_api_key:
            try:
                self.llamaparse = LlamaParse(
                    api_key=llamaparse_api_key,
                    result_type="markdown",  # Get structured markdown output
                    verbose=True
                )
                logger.info("✅ LlamaParse initialized for intelligent PDF parsing")
            except Exception as e:
                logger.warning(f"⚠️ LlamaParse initialization failed: {e}, using PyPDF2 fallback")
                self.llamaparse = None
        else:
            logger.info("ℹ️ LlamaParse not configured, using PyPDF2 for PDF parsing")

    async def parse_pdf(self, file_path: Path) -> str:
        """
        Parse PDF file using LlamaParse (preferred) or PyPDF2 (fallback).

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        # Try LlamaParse first (better for complex AAOIFI documents)
        if self.llamaparse:
            try:
                logger.info(f"📄 Parsing PDF with LlamaParse: {file_path.name}")
                documents = await self.llamaparse.aload_data(str(file_path))

                # Combine all document chunks
                text = "\n\n".join([doc.text for doc in documents])
                logger.info(f"✅ LlamaParse extracted {len(text)} characters from {file_path.name}")
                return text
            except Exception as e:
                logger.warning(f"⚠️ LlamaParse failed for {file_path.name}: {e}, falling back to PyPDF2")

        # Fallback to PyPDF2
        try:
            logger.info(f"📄 Parsing PDF with PyPDF2: {file_path.name}")
            reader = PdfReader(str(file_path))
            text = ""

            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                text += f"\n\n--- Page {page_num + 1} ---\n\n{page_text}"

            logger.info(f"✅ PyPDF2 extracted {len(text)} characters from {file_path.name}")
            return text
        except Exception as e:
            logger.error(f"❌ PDF parsing failed for {file_path.name}: {e}")
            raise

    async def parse_docx(self, file_path: Path) -> str:
        """
        Parse DOCX file using python-docx.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            logger.info(f"📄 Parsing DOCX: {file_path.name}")
            doc = Document(str(file_path))

            # Extract all paragraphs
            text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

            logger.info(f"✅ Extracted {len(text)} characters from {file_path.name}")
            return text
        except Exception as e:
            logger.error(f"❌ DOCX parsing failed for {file_path.name}: {e}")
            raise

    async def ingest_document(
        self,
        file_path: Path,
        document_type: str = "aaoifi",
        metadata: Optional[Dict[str, Any]] = None
    ) -> UploadedDocument:
        """
        Parse document and ingest into Graphiti knowledge graph.

        Args:
            file_path: Path to document file
            document_type: Type of document (aaoifi, context, etc.)
            metadata: Additional metadata to store

        Returns:
            UploadedDocument with ingestion details
        """
        try:
            # Determine file type and parse
            suffix = file_path.suffix.lower()

            if suffix == ".pdf":
                text = await self.parse_pdf(file_path)
            elif suffix in [".docx", ".doc"]:
                text = await self.parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")

            # Prepare metadata
            doc_metadata = {
                "filename": file_path.name,
                "file_type": suffix,
                "document_type": document_type,
                "uploaded_at": datetime.now().isoformat(),
                "file_size": file_path.stat().st_size,
                **(metadata or {})
            }

            # Ingest into Graphiti via MCP
            graphiti_mcp_service = get_graphiti_mcp_service()
            episode_id = await graphiti_mcp_service.ingest_document(
                content=text,
                document_name=file_path.name,
                document_type=document_type,
                metadata=doc_metadata
            )

            # Create UploadedDocument response
            uploaded_doc = UploadedDocument(
                id=episode_id,
                filename=file_path.name,
                filesize=file_path.stat().st_size,
                mime_type=doc_metadata.get("content_type", "application/pdf"),
                uploaded_at=datetime.now(),
                parsed=True,
                episode_id=episode_id,
                content_preview=text[:500] + "..." if len(text) > 500 else text  # Preview
            )

            logger.info(f"✅ Document ingested: {file_path.name} → Episode {episode_id}")
            return uploaded_doc

        except Exception as e:
            logger.error(f"❌ Document ingestion failed for {file_path.name}: {e}")
            raise

    async def save_upload(self, file_content: bytes, filename: str) -> Path:
        """
        Save uploaded file to disk.

        Args:
            file_content: File content bytes
            filename: Original filename

        Returns:
            Path to saved file
        """
        file_path = self.upload_dir / filename

        async with aiofiles.open(file_path, "wb") as f:
            await f.write(file_content)

        logger.info(f"💾 Saved upload: {filename} ({len(file_content)} bytes)")
        return file_path

    async def generate_document(
        self,
        content: str,
        format: str = "pdf",
        filename: Optional[str] = None
    ) -> Path:
        """
        Generate document in specified format using Pandoc.

        Args:
            content: Markdown content to convert
            format: Output format (pdf, docx, markdown)
            filename: Optional output filename

        Returns:
            Path to generated document
        """
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"output_{timestamp}.{format}"

            output_path = self.output_dir / filename

            if format == "markdown":
                # Save markdown directly
                async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                    await f.write(content)
            else:
                # Use Pandoc for PDF/DOCX
                extra_args = []
                if format == "pdf":
                    extra_args = [
                        "--pdf-engine=xelatex",
                        "-V", "geometry:margin=1in"
                    ]

                pypandoc.convert_text(
                    content,
                    format,
                    format="markdown",
                    outputfile=str(output_path),
                    extra_args=extra_args
                )

            logger.info(f"📄 Generated {format.upper()} document: {filename}")
            return output_path

        except Exception as e:
            logger.error(f"❌ Document generation failed: {e}")
            raise


# Singleton instance
_document_service: Optional[DocumentService] = None


def get_document_service() -> DocumentService:
    """Get or create DocumentService singleton."""
    global _document_service

    if _document_service is None:
        llamaparse_api_key = os.getenv("LLAMAPARSE_API_KEY")
        upload_dir = os.getenv("UPLOAD_DIR", "./uploads")
        output_dir = os.getenv("OUTPUT_DIR", "./outputs")

        _document_service = DocumentService(
            llamaparse_api_key=llamaparse_api_key,
            upload_dir=upload_dir,
            output_dir=output_dir
        )

    return _document_service
